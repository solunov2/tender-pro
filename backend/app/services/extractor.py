"""
Tender AI Platform - Document Text Extraction Service
Supports PDF (digital/scanned), DOCX, XLSX - all in-memory

WORKFLOW:
1. First-page scan of ALL files → classify to find Avis
2. Full extraction of Avis ONLY
3. Store Avis text → Run AI pipeline
"""

import io
import re
from typing import Dict, Optional, Tuple, List
from dataclasses import dataclass
from enum import Enum
from loguru import logger

# Document processing
import pypdf
from docx import Document as DocxDocument
import openpyxl
import pandas as pd


class DocumentType(str, Enum):
    AVIS = "AVIS"
    RC = "RC"
    CPS = "CPS"
    ANNEXE = "ANNEXE"
    BPDE = "BPDE"  # Bordereau des Prix - Détail Estimatif
    AE = "AE"  # Acte d'Engagement
    DSH = "DSH"  # Décomposition du montant global / Sous-détail
    CCAG = "CCAG"  # Cahier des Clauses Administratives Générales
    CCTP = "CCTP"  # Cahier des Clauses Techniques Particulières
    BQ = "BQ"  # Bordereau des Quantités
    DQE = "DQE"  # Devis Quantitatif Estimatif
    OTHER = "OTHER"  # Identified but not in main categories
    UNKNOWN = "UNKNOWN"


class ExtractionMethod(str, Enum):
    DIGITAL = "DIGITAL"
    OCR = "OCR"


@dataclass
class FirstPageResult:
    """Result of first-page scan for classification"""
    filename: str
    first_page_text: str
    document_type: DocumentType
    is_scanned: bool
    mime_type: str
    file_size_bytes: int
    success: bool
    error: Optional[str] = None


@dataclass
class ExtractionResult:
    """Result of full text extraction from a document"""
    filename: str
    document_type: DocumentType
    text: str
    page_count: Optional[int]
    extraction_method: ExtractionMethod
    file_size_bytes: int
    mime_type: str
    success: bool
    error: Optional[str] = None


# Classification keywords for document type detection
# Priority order matters - AVIS is checked FIRST
CLASSIFICATION_KEYWORDS = {
    DocumentType.AVIS: [
        "avis de consultation",
        "avis d'appel d'offres", 
        "avis d'appel",
        "avis appel offres",
        "avis ao",
        "avis",  # Generic - filename must contain "avis"
    ],
    DocumentType.RC: [
        "règlement de consultation",
        "reglement de consultation",
        "règlement de la consultation",
        "reglement de la consultation",
    ],
    DocumentType.CPS: [
        "cahier des prescriptions spéciales",
        "cahier des prescriptions speciales",
        "cahier des clauses",
    ],
    DocumentType.ANNEXE: [
        "annexe",
        "additif",
        "avenant",
    ]
}

# Filename patterns for classification (regex patterns)
FILENAME_PATTERNS = {
    DocumentType.AVIS: [
        r'\bavis\b',           # "avis" as whole word
        r'\bavis[\s_-]',       # "avis " or "avis_" or "avis-"
        r'[\s_-]avis\b',       # " avis" or "_avis" or "-avis"
        r'avis[\s_-]*(ar|fr)', # "avis ar" or "avis fr" (Arabic/French)
    ],
    DocumentType.RC: [
        r'\brc\b',             # "rc" as whole word
        r'\brcdp\b',           # "rcdp"
        r'\brcdg\b',           # "rcdg"
    ],
    DocumentType.CPS: [
        r'\bcps\b',            # "cps" as whole word
        r'\bccaf\b',           # "ccaf"
    ],
    DocumentType.CCTP: [
        r'\bcctp\b',           # "cctp" (cahier des clauses techniques)
    ],
    DocumentType.ANNEXE: [
        r'\bannexe\b',
    ],
    DocumentType.BPDE: [
        r'\bbpde\b',           # "bpde"
        r'\bbordereau[\s_-]*prix\b',  # "bordereau prix"
        r'\bbdp\b',            # "bdp"
    ],
    DocumentType.AE: [
        r'\bae\b',             # "ae" as whole word
        r'\bacte[\s_-]*engagement\b',  # "acte engagement"
    ],
    DocumentType.DSH: [
        r'\bdsh\b',            # "dsh"
        r'\bsous[\s_-]*detail\b',  # "sous detail"
        r'\bdecomposition\b',  # "decomposition"
    ],
    DocumentType.CCAG: [
        r'\bccag\b',           # "ccag"
    ],
    DocumentType.BQ: [
        r'\bbq\b',             # "bq"
        r'\bbordereau[\s_-]*quantit\b',  # "bordereau quantit..."
    ],
    DocumentType.DQE: [
        r'\bdqe\b',            # "dqe"
        r'\bdevis[\s_-]*quantitatif\b',  # "devis quantitatif"
    ],
}


def classify_document(text: str, filename: str = "", use_ai: bool = False, is_scanned: bool = False) -> DocumentType:
    """
    Classify document type by scanning first-page content and filename.
    Priority: AVIS > RC > CPS > other document types
    
    Args:
        text: First page text content
        filename: Document filename
        use_ai: Whether to use AI classification as fallback
        is_scanned: Whether document is scanned (limits text for AI)
    """
    text_lower = text.lower()
    filename_lower = filename.lower()
    
    # Extract just the file name without path
    base_filename = filename_lower.split('/')[-1].split('\\')[-1]
    
    # All document types to check
    all_doc_types = [
        DocumentType.AVIS, DocumentType.RC, DocumentType.CPS, DocumentType.ANNEXE,
        DocumentType.BPDE, DocumentType.AE, DocumentType.DSH, DocumentType.CCAG,
        DocumentType.CCTP, DocumentType.BQ, DocumentType.DQE
    ]
    
    # PRIORITY 1: Check filename patterns (most reliable)
    for doc_type in all_doc_types:
        if doc_type in FILENAME_PATTERNS:
            for pattern in FILENAME_PATTERNS[doc_type]:
                if re.search(pattern, base_filename, re.IGNORECASE):
                    # For AVIS, make sure it's not RC/CPS file with "avis" in name
                    if doc_type == DocumentType.AVIS:
                        # Exclude if filename clearly indicates RC or CPS
                        if re.search(r'\b(rc|cps|ccaf|rcdp|rcdg)\b', base_filename):
                            continue
                    return doc_type
    
    # PRIORITY 2: Check text content keywords
    for doc_type in [DocumentType.AVIS, DocumentType.RC, DocumentType.CPS, DocumentType.ANNEXE]:
        if doc_type in CLASSIFICATION_KEYWORDS:
            for keyword in CLASSIFICATION_KEYWORDS[doc_type]:
                if keyword in text_lower:
                    return doc_type
    
    # PRIORITY 3: Use AI classification if enabled and text available
    if use_ai and text and len(text.strip()) > 20:
        ai_result = classify_document_with_ai(text, filename, is_scanned)
        if ai_result != DocumentType.UNKNOWN:
            return ai_result
    
    return DocumentType.UNKNOWN


def classify_document_with_ai(text: str, filename: str = "", is_scanned: bool = False) -> DocumentType:
    """
    Use DeepSeek AI to classify document type.
    For scanned documents, limits to first 500 words.
    
    Args:
        text: Document text content
        filename: Document filename
        is_scanned: Whether document is scanned (OCR'd)
    
    Returns:
        DocumentType classification
    """
    try:
        from openai import OpenAI
        from app.core.config import settings
        
        if not settings.DEEPSEEK_API_KEY:
            logger.warning("DeepSeek API key not configured, skipping AI classification")
            return DocumentType.UNKNOWN
        
        # For scanned docs, use first 500 words only
        if is_scanned:
            words = text.split()[:500]
            text_to_analyze = " ".join(words)
        else:
            # For digital docs, use first 2000 chars
            text_to_analyze = text[:2000]
        
        client = OpenAI(
            api_key=settings.DEEPSEEK_API_KEY,
            base_url=settings.DEEPSEEK_BASE_URL
        )
        
        system_prompt = """You are a document classifier for Moroccan government tender documents (marchés publics).

Classify the document into ONE of these categories:
- AVIS: Avis de consultation, avis d'appel d'offres, notice of tender, announcement
- RC: Règlement de consultation, consultation rules
- CPS: Cahier des prescriptions spéciales, Cahier des charges, specifications document
- ANNEXE: Annexe, addendum, modification document, avenant
- BPDE: Bordereau des Prix - Détail Estimatif, price schedule, pricing breakdown
- AE: Acte d'Engagement, commitment letter, bid submission form
- DSH: Décomposition du montant global, sous-détail des prix, cost breakdown
- CCAG: Cahier des Clauses Administratives Générales
- CCTP: Cahier des Clauses Techniques Particulières, technical specifications
- BQ: Bordereau des Quantités, quantity schedule
- DQE: Devis Quantitatif Estimatif, estimated quantities and prices
- OTHER: Document identified but doesn't fit above categories

RULES:
1. If document is an announcement/notice for tender → AVIS
2. If document describes rules/procedures for bidders → RC
3. If document contains technical/administrative specifications → CPS or CCTP
4. If document is a pricing form/template → BPDE or DQE
5. If document is a commitment/engagement form → AE
6. If document breaks down costs/prices → DSH
7. If document lists quantities → BQ
8. If you can identify the document type but it doesn't fit categories → OTHER
9. NEVER return UNKNOWN if you can identify any document type

Respond with ONLY one word from the list above."""

        user_content = f"""Filename: {filename}

Document content:
{text_to_analyze}

Classification:"""

        logger.info(f"AI classifying document: {filename}")
        
        response = client.chat.completions.create(
            model=settings.DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            max_tokens=10,
            temperature=0
        )
        
        result = response.choices[0].message.content.strip().upper()
        
        # Map response to DocumentType
        type_map = {
            "AVIS": DocumentType.AVIS,
            "RC": DocumentType.RC,
            "CPS": DocumentType.CPS,
            "ANNEXE": DocumentType.ANNEXE,
            "BPDE": DocumentType.BPDE,
            "AE": DocumentType.AE,
            "DSH": DocumentType.DSH,
            "CCAG": DocumentType.CCAG,
            "CCTP": DocumentType.CCTP,
            "BQ": DocumentType.BQ,
            "DQE": DocumentType.DQE,
            "OTHER": DocumentType.OTHER,
        }
        
        doc_type = type_map.get(result, DocumentType.OTHER)
        logger.info(f"AI classified {filename} as: {doc_type.value}")
        return doc_type
        
    except Exception as e:
        logger.error(f"AI classification failed: {e}")
        return DocumentType.UNKNOWN


# ============================
# FIRST-PAGE EXTRACTION (Classification Phase)
# ============================

def _is_pdf_scanned(file_bytes: io.BytesIO) -> Tuple[bool, str]:
    """
    Check if PDF is scanned by attempting digital extraction of first page.
    Returns (is_scanned, first_page_text)
    """
    file_bytes.seek(0)
    try:
        reader = pypdf.PdfReader(file_bytes)
        if len(reader.pages) == 0:
            return True, ""
        
        first_page_text = reader.pages[0].extract_text() or ""
        
        # If text is too sparse (<100 chars), it's scanned
        is_scanned = len(first_page_text.strip()) < 100
        return is_scanned, first_page_text
    except Exception as e:
        logger.warning(f"PDF scan check failed: {e}")
        return True, ""


def _ocr_first_page_pdf(file_bytes: io.BytesIO) -> str:
    """OCR only the first page of a scanned PDF using Tesseract.
    
    Uses pytesseract with pdf2image for conversion.
    """
    import pytesseract
    from pdf2image import convert_from_bytes
    
    # Configure Tesseract path (Windows)
    TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    POPPLER_PATH = r"C:\poppler-24.08.0\Library\bin"
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    
    try:
        file_bytes.seek(0)
        pdf_bytes = file_bytes.read()
        
        # Convert first page to image
        logger.info("Converting first page to image...")
        images = convert_from_bytes(
            pdf_bytes, 
            dpi=200, 
            first_page=1, 
            last_page=1,
            poppler_path=POPPLER_PATH
        )
        
        if not images:
            logger.error("Could not convert PDF first page to image")
            return ""
        
        # Run Tesseract OCR
        logger.info("Running Tesseract OCR on first page...")
        text = pytesseract.image_to_string(
            images[0], 
            lang="fra+ara+eng",
            config="--oem 3 --psm 3"
        )
        
        logger.info(f"OCR extracted {len(text)} chars from first page")
        return text.strip()
        
    except Exception as e:
        logger.error(f"First-page OCR failed: {e}")
        return ""



def _get_first_page_docx(file_bytes: io.BytesIO) -> str:
    """Get first ~1000 chars from DOCX (approximates first page)"""
    file_bytes.seek(0)
    try:
        doc = DocxDocument(file_bytes)
        text_parts = []
        char_count = 0
        
        for para in doc.paragraphs:
            text_parts.append(para.text)
            char_count += len(para.text)
            if char_count > 1000:
                break
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX first-page extraction failed: {e}")
        return ""


def _get_first_page_doc(file_bytes: io.BytesIO) -> str:
    """
    Extract first page text from legacy .doc files.
    Uses multiple fallback methods.
    """
    file_bytes.seek(0)
    content = file_bytes.read()
    
    # Method 1: Try using antiword via subprocess (if installed)
    try:
        import subprocess
        import tempfile
        
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                import os
                os.unlink(tmp_path)
                # Return first 1000 chars
                return result.stdout[:1000]
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        finally:
            try:
                import os
                os.unlink(tmp_path)
            except:
                pass
    except Exception:
        pass
    
    # Method 2: Basic binary text extraction (fallback)
    try:
        # .doc files often have readable text mixed with binary
        text_parts = []
        # Try to decode as various encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                decoded = content.decode(encoding, errors='ignore')
                # Extract readable text sequences (4+ chars)
                import re
                words = re.findall(r'[a-zA-ZÀ-ÿ\s]{4,}', decoded)
                if words:
                    text = ' '.join(words)
                    # Clean up
                    text = re.sub(r'\s+', ' ', text).strip()
                    if len(text) > 50:  # Reasonable amount of text
                        return text[:1000]
            except:
                continue
    except Exception as e:
        logger.warning(f"Binary .doc extraction failed: {e}")
    
    return ""


def _extract_full_doc(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full extraction from legacy .doc files"""
    file_bytes.seek(0)
    content = file_bytes.read()
    
    # Method 1: Try using antiword via subprocess
    try:
        import subprocess
        import tempfile
        
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                import os
                os.unlink(tmp_path)
                return result.stdout, None
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        finally:
            try:
                import os
                os.unlink(tmp_path)
            except:
                pass
    except Exception:
        pass
    
    # Method 2: Basic binary text extraction
    try:
        text_parts = []
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                decoded = content.decode(encoding, errors='ignore')
                import re
                words = re.findall(r'[a-zA-ZÀ-ÿ0-9\s\.,;:\-\(\)]{4,}', decoded)
                if words:
                    text = ' '.join(words)
                    text = re.sub(r'\s+', ' ', text).strip()
                    if len(text) > 100:
                        return text, None
            except:
                continue
    except Exception as e:
        logger.warning(f"Binary .doc full extraction failed: {e}")
    
    return "[.DOC EXTRACTION FAILED - Install antiword for better support]", None


def _get_first_page_xlsx(file_bytes: io.BytesIO) -> str:
    """Get first rows from first sheet of XLSX"""
    file_bytes.seek(0)
    try:
        wb = openpyxl.load_workbook(file_bytes, read_only=True, data_only=True)
        if not wb.sheetnames:
            wb.close()
            return ""
        
        sheet = wb[wb.sheetnames[0]]
        text_parts = []
        row_count = 0
        
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(cell) if cell is not None else "" for cell in row]
            if any(row_values):
                text_parts.append(" | ".join(row_values))
                row_count += 1
                if row_count > 20:  # First 20 rows
                    break
        
        wb.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"XLSX first-page extraction failed: {e}")
        return ""


def extract_first_page(filename: str, file_bytes: io.BytesIO, use_ai_classification: bool = True) -> FirstPageResult:
    """
    Extract FIRST PAGE ONLY for classification purposes.
    This is a quick scan to identify document type.
    
    Args:
        filename: Document filename
        file_bytes: Document content as BytesIO
        use_ai_classification: Whether to use AI for classification fallback
    """
    # Skip temp files and hidden files
    base_name = filename.split('/')[-1]
    if base_name.startswith('~$') or base_name.startswith('.'):
        return FirstPageResult(
            filename=filename,
            first_page_text="",
            document_type=DocumentType.UNKNOWN,
            is_scanned=False,
            mime_type="",
            file_size_bytes=0,
            success=False,
            error="Temporary or hidden file - skipped"
        )
    
    # Get file size
    file_bytes.seek(0, 2)
    file_size = file_bytes.tell()
    file_bytes.seek(0)
    
    # Determine MIME type from extension
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    mime_map = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls': 'application/vnd.ms-excel',
    }
    mime_type = mime_map.get(ext, 'application/octet-stream')
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    is_scanned = False
    first_page_text = ""
    
    try:
        if ext == 'pdf' or mime_type == 'application/pdf':
            is_scanned, first_page_text = _is_pdf_scanned(file_bytes)
            
            if is_scanned and not first_page_text:
                logger.info(f"Scanned PDF detected, OCR first page: {filename}")
                first_page_text = _ocr_first_page_pdf(file_bytes)
                
        elif ext == 'docx' or 'wordprocessingml' in mime_type:
            first_page_text = _get_first_page_docx(file_bytes)
            
        elif ext == 'doc':
            # Legacy .doc files - try extraction
            logger.info(f"Extracting legacy .doc file: {filename}")
            first_page_text = _get_first_page_doc(file_bytes)
            
        elif ext in ('xls', 'xlsx') or 'excel' in mime_type or 'spreadsheet' in mime_type:
            first_page_text = _get_first_page_xlsx(file_bytes)
            
        elif ext == 'txt' or mime_type == 'text/plain':
            file_bytes.seek(0)
            first_page_text = file_bytes.read(2000).decode('utf-8', errors='ignore')
            
        else:
            return FirstPageResult(
                filename=filename,
                first_page_text="",
                document_type=DocumentType.UNKNOWN,
                is_scanned=False,
                mime_type=mime_type,
                file_size_bytes=file_size,
                success=False,
                error=f"Unsupported file type: {ext}"
            )
        
        # Classify based on first-page content
        # Use AI classification for better accuracy, especially on scanned docs
        doc_type = classify_document(
            first_page_text, 
            filename, 
            use_ai=use_ai_classification,
            is_scanned=is_scanned
        )
        
        return FirstPageResult(
            filename=filename,
            first_page_text=first_page_text,
            document_type=doc_type,
            is_scanned=is_scanned,
            mime_type=mime_type,
            file_size_bytes=file_size,
            success=True
        )
        
    except Exception as e:
        logger.error(f"First-page extraction failed for {filename}: {e}")
        return FirstPageResult(
            filename=filename,
            first_page_text="",
            document_type=DocumentType.UNKNOWN,
            is_scanned=False,
            mime_type=mime_type,
            file_size_bytes=file_size,
            success=False,
            error=str(e)
        )


# ============================
# FULL EXTRACTION (Only for identified Avis)
# ============================

def _extract_full_pdf_digital(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full digital extraction from PDF"""
    file_bytes.seek(0)
    reader = pypdf.PdfReader(file_bytes)
    page_count = len(reader.pages)
    
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text_parts.append(page_text)
    
    return "\n\n".join(text_parts), page_count


def _extract_full_pdf_ocr(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full OCR extraction from scanned PDF using Tesseract.
    
    Uses pytesseract with pdf2image for conversion.
    """
    import pytesseract
    from pdf2image import convert_from_bytes
    
    # Configure Tesseract path (Windows)
    TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    POPPLER_PATH = r"C:\poppler-24.08.0\Library\bin"
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    
    try:
        logger.info("Full OCR extraction starting (Tesseract)...")
        
        file_bytes.seek(0)
        pdf_bytes = file_bytes.read()
        
        # Convert all pages to images
        logger.info("Converting PDF to images...")
        images = convert_from_bytes(
            pdf_bytes, 
            dpi=200,
            poppler_path=POPPLER_PATH
        )
        
        if not images:
            logger.error("Could not convert PDF to images")
            return "[OCR FAILED: No images extracted]", 0
        
        logger.info(f"Converted {len(images)} pages, running Tesseract OCR...")
        
        all_text = []
        for i, image in enumerate(images):
            logger.info(f"OCR page {i + 1}/{len(images)}...")
            page_text = pytesseract.image_to_string(
                image, 
                lang="fra+ara+eng",
                config="--oem 3 --psm 3"
            )
            all_text.append(f"--- Page {i + 1} ---\n{page_text}")
        
        logger.info(f"OCR completed: {len(images)} pages")
        return "\n\n".join(all_text).strip(), len(images)
        
    except Exception as e:
        logger.error(f"Full OCR failed: {e}")
        return f"[OCR FAILED: {str(e)}]", 0



def _extract_full_docx(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full extraction from DOCX"""
    file_bytes.seek(0)
    doc = DocxDocument(file_bytes)
    
    paragraphs = [para.text for para in doc.paragraphs]
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            paragraphs.append(row_text)
    
    return "\n".join(paragraphs), None


def _extract_full_xlsx(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full extraction from Excel"""
    file_bytes.seek(0)
    
    try:
        wb = openpyxl.load_workbook(file_bytes, read_only=True, data_only=True)
        
        all_text = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            all_text.append(f"=== Sheet: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(row_values):
                    all_text.append(" | ".join(row_values))
        
        wb.close()
        return "\n".join(all_text), None
        
    except Exception:
        file_bytes.seek(0)
        try:
            df = pd.read_excel(file_bytes, sheet_name=None)
            all_text = []
            for sheet_name, sheet_df in df.items():
                all_text.append(f"=== Sheet: {sheet_name} ===")
                all_text.append(sheet_df.to_string())
            return "\n".join(all_text), None
        except Exception as e:
            return f"[EXCEL EXTRACTION FAILED: {e}]", None


def extract_full_document(filename: str, file_bytes: io.BytesIO, is_scanned: bool = False) -> ExtractionResult:
    """
    Full extraction of a single document.
    Use appropriate method based on is_scanned flag.
    """
    file_bytes.seek(0, 2)
    file_size = file_bytes.tell()
    file_bytes.seek(0)
    
    # Determine MIME type from extension
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    mime_map = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls': 'application/vnd.ms-excel',
        'txt': 'text/plain',
    }
    mime_type = mime_map.get(ext, 'application/octet-stream')
    
    try:
        if ext == 'pdf' or mime_type == 'application/pdf':
            if is_scanned:
                text, page_count = _extract_full_pdf_ocr(file_bytes)
                method = ExtractionMethod.OCR
            else:
                text, page_count = _extract_full_pdf_digital(file_bytes)
                method = ExtractionMethod.DIGITAL
                
        elif ext == 'docx' or 'wordprocessingml' in mime_type:
            text, page_count = _extract_full_docx(file_bytes)
            method = ExtractionMethod.DIGITAL
            
        elif ext == 'doc':
            text, page_count = _extract_full_doc(file_bytes)
            method = ExtractionMethod.DIGITAL
            
        elif ext in ('xls', 'xlsx') or 'excel' in mime_type or 'spreadsheet' in mime_type:
            text, page_count = _extract_full_xlsx(file_bytes)
            method = ExtractionMethod.DIGITAL
            
        elif ext == 'txt' or mime_type == 'text/plain':
            file_bytes.seek(0)
            text = file_bytes.read().decode('utf-8', errors='ignore')
            page_count = None
            method = ExtractionMethod.DIGITAL
            
        else:
            return ExtractionResult(
                filename=filename,
                document_type=DocumentType.UNKNOWN,
                text="",
                page_count=None,
                extraction_method=ExtractionMethod.DIGITAL,
                file_size_bytes=file_size,
                mime_type=mime_type,
                success=False,
                error=f"Unsupported file type: {ext}"
            )
        
        doc_type = classify_document(text, filename)
        
        return ExtractionResult(
            filename=filename,
            document_type=doc_type,
            text=text,
            page_count=page_count,
            extraction_method=method,
            file_size_bytes=file_size,
            mime_type=mime_type,
            success=True
        )
        
    except Exception as e:
        logger.error(f"Full extraction failed for {filename}: {e}")
        return ExtractionResult(
            filename=filename,
            document_type=DocumentType.UNKNOWN,
            text="",
            page_count=None,
            extraction_method=ExtractionMethod.DIGITAL,
            file_size_bytes=file_size,
            mime_type=mime_type,
            success=False,
            error=str(e)
        )


# ============================
# MAIN WORKFLOW: Classify then Extract Avis
# ============================

def classify_all_documents(zip_files: Dict[str, io.BytesIO]) -> List[FirstPageResult]:
    """
    STEP 1: Scan first page of ALL files to classify them.
    Returns classification results (first-page text is temporary, will be discarded).
    """
    results = []
    
    for filename, file_bytes in zip_files.items():
        # Skip hidden files and directories
        if filename.startswith('.') or filename.startswith('__'):
            continue
        
        logger.info(f"Classifying: {filename}")
        result = extract_first_page(filename, file_bytes)
        results.append(result)
    
    return results


def _is_french_document(filename: str, first_page_text: str) -> bool:
    """
    Check if document is French version based on filename and content.
    Uses strict patterns to avoid false positives.
    """
    filename_lower = filename.lower()
    text_lower = first_page_text.lower() if first_page_text else ""
    
    # Strict filename patterns for French (word boundaries)
    french_filename_patterns = [
        r'[\s_\-\.]fr[\s_\-\.]',      # _fr_ or -fr- or .fr.
        r'[\s_\-\.]fr$',              # ends with _fr or -fr
        r'^fr[\s_\-\.]',              # starts with fr_ or fr-
        r'[\s_\-]français',           # _français
        r'[\s_\-]francais',           # _francais
        r'[\s_\-]french',             # _french
        r'\(fr\)',                    # (fr)
        r'\[fr\]',                    # [fr]
        r'version[\s_\-]*fr',         # version fr
    ]
    
    for pattern in french_filename_patterns:
        if re.search(pattern, filename_lower):
            return True
    
    # Check content for French language indicators
    french_content_markers = [
        'règlement de consultation',
        'cahier des prescriptions',
        'avis d\'appel d\'offres',
        'marché public',
        'le soumissionnaire',
        'pièces justificatives',
    ]
    
    if text_lower:
        french_score = sum(1 for marker in french_content_markers if marker in text_lower)
        if french_score >= 2:
            return True
    
    return False


def _is_arabic_document(filename: str, first_page_text: str) -> bool:
    """
    Check if document is Arabic version based on filename and content.
    Uses strict patterns to avoid false positives.
    """
    filename_lower = filename.lower()
    text_lower = first_page_text.lower() if first_page_text else ""
    
    # Strict filename patterns for Arabic (word boundaries)
    arabic_filename_patterns = [
        r'[\s_\-\.]ar[\s_\-\.]',      # _ar_ or -ar- or .ar.
        r'[\s_\-\.]ar$',              # ends with _ar or -ar
        r'^ar[\s_\-\.]',              # starts with ar_ or ar-
        r'[\s_\-]arabe',              # _arabe
        r'[\s_\-]arabic',             # _arabic
        r'\(ar\)',                    # (ar)
        r'\[ar\]',                    # [ar]
        r'version[\s_\-]*ar',         # version ar
        r'عربي',                      # Arabic word for "Arabic"
        r'العربية',                   # Arabic for "Arabic language"
    ]
    
    for pattern in arabic_filename_patterns:
        if re.search(pattern, filename_lower):
            return True
    
    # Check for Arabic script in content (significant presence)
    if first_page_text:
        arabic_chars = len(re.findall(r'[\u0600-\u06FF]', first_page_text))
        latin_chars = len(re.findall(r'[a-zA-Z]', first_page_text))
        # If Arabic chars dominate, it's an Arabic document
        if arabic_chars > latin_chars and arabic_chars > 50:
            return True
    
    return False


def _is_multi_tender_avis(first_page_text: str, tender_reference: Optional[str] = None) -> bool:
    """
    Detect if an Avis contains multiple tenders (compiled from same organization).
    
    Indicators of multi-tender Avis:
    1. Multiple reference numbers in the first page
    2. Phrases like "appels d'offres suivants", "marchés suivants"
    3. Table-like structure with multiple tender entries
    """
    if not first_page_text:
        return False
    
    text_lower = first_page_text.lower()
    
    # Check for multi-tender phrases
    multi_tender_indicators = [
        "appels d'offres suivants",
        "marchés suivants",
        "consultations suivantes",
        "liste des appels",
        "tableau des marchés",
        "les références ci-après",
        "les marchés ci-après",
    ]
    
    for indicator in multi_tender_indicators:
        if indicator in text_lower:
            logger.warning(f"Multi-tender indicator found: '{indicator}'")
            return True
    
    # Count reference patterns (e.g., "N° 01/2024", "Ref: 123/2024")
    ref_patterns = [
        r'n[°o]?\s*\d+[/\-]\d{4}',
        r'ref[:\s]+\d+[/\-]\d{4}',
        r'\d+[/\-]ao[/\-]\d{4}',
    ]
    
    ref_count = 0
    for pattern in ref_patterns:
        matches = re.findall(pattern, text_lower)
        ref_count += len(matches)
    
    # If more than 3 reference-like patterns, likely multi-tender
    if ref_count > 3:
        logger.warning(f"Multiple reference patterns detected ({ref_count}), likely multi-tender Avis")
        return True
    
    return False


def _select_best_document(candidates: List[FirstPageResult], doc_type_name: str) -> Optional[FirstPageResult]:
    """
    Select best document from candidates, prioritizing French over Arabic.
    """
    if not candidates:
        return None
    
    french_docs = []
    arabic_docs = []
    neutral_docs = []
    
    for doc in candidates:
        is_french = _is_french_document(doc.filename, doc.first_page_text)
        is_arabic = _is_arabic_document(doc.filename, doc.first_page_text)
        
        logger.debug(f"  {doc.filename}: french={is_french}, arabic={is_arabic}")
        
        if is_french and not is_arabic:
            french_docs.append(doc)
        elif is_arabic and not is_french:
            arabic_docs.append(doc)
        else:
            neutral_docs.append(doc)
    
    # Priority 1: Explicit French documents
    if french_docs:
        logger.success(f"Selected French {doc_type_name}: {french_docs[0].filename}")
        return french_docs[0]
    
    # Priority 2: Neutral documents (not marked as Arabic)
    if neutral_docs:
        logger.success(f"Selected neutral {doc_type_name} (not Arabic): {neutral_docs[0].filename}")
        return neutral_docs[0]
    
    # Priority 3: Arabic documents (only if no other choice)
    if arabic_docs:
        logger.warning(f"Only Arabic {doc_type_name} available: {arabic_docs[0].filename}")
        return arabic_docs[0]
    
    return candidates[0]


def find_primary_document(
    classifications: List[FirstPageResult],
    tender_reference: Optional[str] = None
) -> Tuple[Optional[FirstPageResult], str]:
    """
    STEP 2: Find the primary document for extraction.
    
    PRIORITY:
    1. Avis de Consultation (if valid, single-tender)
    2. CPS (fallback if Avis not found OR Avis is multi-tender)
    
    Returns:
        Tuple of (document_info, source_type) where source_type is "AVIS" or "CPS"
    """
    avis_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.AVIS]
    cps_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.CPS]
    
    logger.info(f"Found {len(avis_candidates)} Avis, {len(cps_candidates)} CPS candidates")
    
    # Try Avis first
    if avis_candidates:
        logger.info(f"Evaluating {len(avis_candidates)} Avis candidates...")
        best_avis = _select_best_document(avis_candidates, "Avis")
        
        if best_avis:
            # Check if this Avis is a multi-tender compilation
            if _is_multi_tender_avis(best_avis.first_page_text, tender_reference):
                logger.warning(f"Avis '{best_avis.filename}' appears to be multi-tender compilation")
                logger.info("Falling back to CPS...")
            else:
                logger.success(f"Using Avis: {best_avis.filename}")
                return best_avis, "AVIS"
    else:
        logger.warning("No Avis document found in ZIP")
    
    # Fallback to CPS
    if cps_candidates:
        logger.info(f"Falling back to CPS ({len(cps_candidates)} candidates)...")
        best_cps = _select_best_document(cps_candidates, "CPS")
        
        if best_cps:
            logger.success(f"Using CPS as fallback: {best_cps.filename}")
            return best_cps, "CPS"
    
    logger.error("No Avis or CPS found - cannot extract tender metadata")
    return None, ""


def find_avis_document(classifications: List[FirstPageResult]) -> Optional[FirstPageResult]:
    """
    STEP 2: Find the Avis de Consultation from classification results.
    Returns the Avis document info, or None if not found.
    
    LEGACY WRAPPER: Now uses find_primary_document internally.
    For new code, use find_primary_document which also returns source type.
    
    PRIORITY: 
    1. Avis (if valid single-tender)
    2. CPS (fallback if Avis not found or is multi-tender)
    """
    doc, source_type = find_primary_document(classifications)
    return doc


def extract_avis_only(
    zip_files: Dict[str, io.BytesIO],
    avis_info: FirstPageResult
) -> Optional[ExtractionResult]:
    """
    STEP 3: Extract FULL content of Avis document only.
    Uses appropriate method (digital/OCR) based on classification.
    """
    if avis_info.filename not in zip_files:
        logger.error(f"Avis file not found in ZIP: {avis_info.filename}")
        return None
    
    file_bytes = zip_files[avis_info.filename]
    
    logger.info(f"Full extraction of Avis: {avis_info.filename} (scanned={avis_info.is_scanned})")
    return extract_full_document(avis_info.filename, file_bytes, avis_info.is_scanned)


def process_tender_zip(
    zip_files: Dict[str, io.BytesIO],
    tender_reference: Optional[str] = None
) -> Tuple[Optional[ExtractionResult], List[FirstPageResult], str]:
    """
    MAIN WORKFLOW: Process a tender ZIP file.
    
    1. Classify all documents (first-page scan)
    2. Find primary document (Avis preferred, CPS as fallback)
    3. Extract full content of primary document
    4. Return (extraction_result, all_classifications, source_type)
    
    Args:
        zip_files: Dictionary of filename -> file bytes
        tender_reference: Optional tender reference to help detect multi-tender Avis
    
    Returns:
        Tuple of (extraction_result, classifications, source_type)
        source_type is "AVIS" or "CPS"
    
    Classifications are returned for logging/debugging but their first_page_text
    should be discarded after processing.
    """
    # Step 1: Classify all documents
    logger.info("Phase 1: Classifying all documents...")
    classifications = classify_all_documents(zip_files)
    
    # Log classification results
    for c in classifications:
        status = "✓" if c.success else "✗"
        scanned = " [SCANNED]" if c.is_scanned else ""
        logger.info(f"  {status} {c.filename} → {c.document_type.value}{scanned}")
    
    # Step 2: Find primary document (Avis or CPS fallback)
    logger.info("Phase 2: Locating primary document (Avis or CPS)...")
    primary_doc, source_type = find_primary_document(classifications, tender_reference)
    
    if not primary_doc:
        return None, classifications, ""
    
    # Step 3: Full extraction of primary document
    logger.info(f"Phase 3: Full extraction of {source_type}...")
    
    if primary_doc.filename not in zip_files:
        logger.error(f"Primary document not found in ZIP: {primary_doc.filename}")
        return None, classifications, ""
    
    file_bytes = zip_files[primary_doc.filename]
    extraction = extract_full_document(primary_doc.filename, file_bytes, primary_doc.is_scanned)
    
    # Update document type based on source
    if extraction and source_type == "CPS":
        extraction.document_type = DocumentType.CPS
    
    # Clear first-page texts from memory (they're no longer needed)
    for c in classifications:
        c.first_page_text = ""  # Discard
    
    return extraction, classifications, source_type


def extract_best_documents_for_phase1(
    zip_files: Dict[str, io.BytesIO],
    tender_reference: Optional[str] = None,
) -> Tuple[Dict[DocumentType, ExtractionResult], List[FirstPageResult]]:
    """Select and fully extract the best AVIS/RC/CPS documents for Phase-1 fallback logic.

    AVIS is ignored if it looks like a multi-tender compilation.

    Returns:
        (extractions_by_type, classifications)
    """
    logger.info("Phase 1: Classifying all documents (for fallbacks)...")
    classifications = classify_all_documents(zip_files)

    avis_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.AVIS]
    rc_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.RC]
    cps_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.CPS]

    best_avis = _select_best_document(avis_candidates, "Avis") if avis_candidates else None
    if best_avis and _is_multi_tender_avis(best_avis.first_page_text, tender_reference):
        logger.warning(f"Ignoring Avis '{best_avis.filename}' (multi-tender compilation)")
        best_avis = None

    best_rc = _select_best_document(rc_candidates, "RC") if rc_candidates else None
    best_cps = _select_best_document(cps_candidates, "CPS") if cps_candidates else None

    selections = [best_avis, best_rc, best_cps]
    extractions: Dict[DocumentType, ExtractionResult] = {}

    for info in selections:
        if not info:
            continue
        if info.filename not in zip_files:
            continue

        file_bytes = zip_files[info.filename]
        logger.info(f"Full extraction of {info.document_type.value}: {info.filename} (scanned={info.is_scanned})")
        extracted = extract_full_document(info.filename, file_bytes, info.is_scanned)
        if extracted and extracted.success:
            extracted.document_type = info.document_type
            extractions[info.document_type] = extracted

    # Discard first-page texts
    for c in classifications:
        c.first_page_text = ""

    return extractions, classifications


def extract_best_documents_for_phase1_lazy(
    zip_files: Dict[str, io.BytesIO],
    tender_reference: Optional[str] = None,
    current_metadata: Optional[Dict] = None,
) -> Tuple[Dict[DocumentType, ExtractionResult], List["FirstPageResult"]]:
    """LAZY EXTRACTION: Only extract/OCR documents when needed for missing fields.
    
    Workflow:
    1. Classify all documents (first-page scan - fast, no full OCR)
    2. For each document type in priority order (AVIS → RC → CPS):
       - Check if we still have missing fields
       - Only then do full extraction (with OCR if scanned)
       - Stop when all fields are satisfied
    
    Args:
        zip_files: Dictionary of filename -> file bytes
        tender_reference: Optional tender reference to detect multi-tender Avis
        current_metadata: Current metadata state (to know what's missing)
    
    Returns:
        (extractions_by_type, classifications)
    """
    from app.services.phase1_merge import is_metadata_complete, get_missing_fields
    
    logger.info("Phase 1 (LAZY): Classifying all documents...")
    classifications = classify_all_documents(zip_files)

    avis_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.AVIS]
    rc_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.RC]
    cps_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.CPS]

    best_avis = _select_best_document(avis_candidates, "Avis") if avis_candidates else None
    if best_avis and _is_multi_tender_avis(best_avis.first_page_text, tender_reference):
        logger.warning(f"Ignoring Avis '{best_avis.filename}' (multi-tender compilation)")
        best_avis = None

    best_rc = _select_best_document(rc_candidates, "RC") if rc_candidates else None
    best_cps = _select_best_document(cps_candidates, "CPS") if cps_candidates else None

    # Ordered by fallback priority
    candidates = [
        (DocumentType.AVIS, best_avis),
        (DocumentType.RC, best_rc),
        (DocumentType.CPS, best_cps),
    ]
    
    extractions: Dict[DocumentType, ExtractionResult] = {}

    for doc_type, info in candidates:
        if not info:
            continue
        if info.filename not in zip_files:
            continue
            
        # Check if we still need this extraction
        if is_metadata_complete(current_metadata):
            logger.info(f"Metadata complete, skipping {doc_type.value} extraction")
            break
        
        missing = get_missing_fields(current_metadata)
        logger.info(f"Missing fields: {missing}. Extracting {doc_type.value}...")

        file_bytes = zip_files[info.filename]
        logger.info(f"Full extraction of {doc_type.value}: {info.filename} (scanned={info.is_scanned})")
        extracted = extract_full_document(info.filename, file_bytes, info.is_scanned)
        
        if extracted and extracted.success:
            extracted.document_type = doc_type
            extractions[doc_type] = extracted

    # Discard first-page texts
    for c in classifications:
        c.first_page_text = ""

    return extractions, classifications


def extract_all_documents_for_phase2(
    zip_files: Dict[str, io.BytesIO],
    tender_reference: Optional[str] = None,
) -> Tuple[Dict[DocumentType, ExtractionResult], List["FirstPageResult"]]:
    """Extract ALL documents (AVIS, RC, CPS, ANNEXE) for Phase 2 deep analysis.
    
    This function extracts all available documents regardless of Phase 1 completion,
    as Phase 2 needs access to all documents for comprehensive analysis.
    
    Args:
        zip_files: Dictionary of filename -> file bytes
        tender_reference: Optional tender reference to detect multi-tender Avis
    
    Returns:
        (extractions_by_type, classifications)
    """
    logger.info("Phase 2: Extracting all documents for deep analysis...")
    classifications = classify_all_documents(zip_files)

    avis_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.AVIS]
    rc_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.RC]
    cps_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.CPS]
    annexe_candidates = [r for r in classifications if r.success and r.document_type == DocumentType.ANNEXE]

    best_avis = _select_best_document(avis_candidates, "Avis") if avis_candidates else None
    if best_avis and _is_multi_tender_avis(best_avis.first_page_text, tender_reference):
        logger.warning(f"Ignoring Avis '{best_avis.filename}' (multi-tender compilation)")
        best_avis = None

    best_rc = _select_best_document(rc_candidates, "RC") if rc_candidates else None
    best_cps = _select_best_document(cps_candidates, "CPS") if cps_candidates else None
    best_annexe = _select_best_document(annexe_candidates, "Annexe") if annexe_candidates else None

    # Extract all documents (priority order: ANNEXE > CPS > RC > AVIS)
    candidates = [
        (DocumentType.ANNEXE, best_annexe),
        (DocumentType.CPS, best_cps),
        (DocumentType.RC, best_rc),
        (DocumentType.AVIS, best_avis),
    ]
    
    extractions: Dict[DocumentType, ExtractionResult] = {}

    for doc_type, info in candidates:
        if not info:
            continue
        if info.filename not in zip_files:
            continue
        
        file_bytes = zip_files[info.filename]
        logger.info(f"Extracting {doc_type.value} for Phase 2: {info.filename} (scanned={info.is_scanned})")
        extracted = extract_full_document(info.filename, file_bytes, info.is_scanned)
        
        if extracted and extracted.success:
            extracted.document_type = doc_type
            extractions[doc_type] = extracted

    # Discard first-page texts
    for c in classifications:
        c.first_page_text = ""

    return extractions, classifications


# ============================
# LEGACY FUNCTION (for backward compatibility)
# ============================

def extract_all_from_zip(zip_files: Dict[str, io.BytesIO]) -> List[ExtractionResult]:
    """
    DEPRECATED: Old function that extracted all files fully.
    Kept for backward compatibility but should use process_tender_zip instead.
    """
    logger.warning("Using legacy extract_all_from_zip - consider using process_tender_zip")
    
    results = []
    for filename, file_bytes in zip_files.items():
        if filename.startswith('.') or filename.startswith('__'):
            continue
        
        # Check if scanned first
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        is_scanned = False
        if ext == 'pdf':
            is_scanned, _ = _is_pdf_scanned(file_bytes)
        
        result = extract_full_document(filename, file_bytes, is_scanned)
        results.append(result)
    
    return results
